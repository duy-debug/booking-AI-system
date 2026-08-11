from dataclasses import dataclass
from pathlib import Path

# ============================================================
# THÊM MỚI:
# Import thư viện để đọc PDF và DOCX
# ============================================================

from pypdf import PdfReader

# python-docx cũng có class tên Document.
# Vì code của chúng ta cũng đang có dataclass tên Document,
# nếu import kiểu:
#
# from docx import Document
#
# thì sẽ bị trùng tên.
#
# Vì vậy ta alias thành DocxDocument.
from docx import Document as DocxDocument


# ============================================================
# Document
# ============================================================
#
# Đây là cấu trúc dữ liệu đại diện cho MỘT tài liệu sau khi load.
#
# Ví dụ:
#
# booking_policy.md
#        ↓
#
# Document(
#     text="Nội dung chính sách...",
#     source="booking_policy.md",
#     file_path="knowledge/booking_policy.md"
# )
#
# Ở bước loader:
#
# - text      = toàn bộ nội dung file sau khi extract
# - source    = tên file
# - file_path = đường dẫn file
#
# Sau này chunker.py sẽ nhận Document này và chia nó
# thành nhiều chunk nhỏ hơn.
# ============================================================

@dataclass
class Document:
    text: str
    source: str
    file_path: str


# ============================================================
# DocumentLoader
# ============================================================
#
# Class này chịu trách nhiệm:
#
# File / Folder
#       ↓
# xác định loại file
#       ↓
# dùng parser phù hợp
#       ↓
# lấy text
#       ↓
# Document
#
# Nó KHÔNG làm:
#
# - chunking
# - embedding
# - lưu Qdrant
# - retrieval
# ============================================================

class DocumentLoader:

    # ========================================================
    # SỬA CHỖ 1:
    # Mở rộng danh sách file hỗ trợ.
    # ========================================================
    #
    # Trước:
    #
    # SUPPORTED_EXTENSIONS = {".md", ".txt"}
    #
    # Sau:
    #
    # - .md   → đọc bằng read_text()
    # - .txt  → đọc bằng read_text()
    # - .pdf  → đọc bằng pypdf
    # - .docx → đọc bằng python-docx
    #
    # Lưu ý:
    #
    # .doc và .docx KHÔNG giống nhau.
    #
    # python-docx hỗ trợ .docx,
    # không trực tiếp hỗ trợ file Word cũ .doc.
    # ========================================================

    SUPPORTED_EXTENSIONS = {
        ".md",
        ".txt",
        ".pdf",
        ".docx",
    }


    def load_file(
        self,
        file_path: str | Path,
    ) -> Document:
        """
        Load một file knowledge và chuyển thành Document.

        Flow:

        file_path
            ↓
        kiểm tra file tồn tại
            ↓
        kiểm tra extension
            ↓
        chọn parser phù hợp
            ↓
        extract text
            ↓
        tạo Document
            ↓
        return
        """

        # ----------------------------------------------------
        # 1. Chuyển input thành Path object
        # ----------------------------------------------------

        path = Path(file_path)


        # ----------------------------------------------------
        # 2. Kiểm tra file có tồn tại không
        # ----------------------------------------------------

        if not path.exists():
            raise FileNotFoundError(
                f"Knowledge file does not exist: {path}"
            )


        # ----------------------------------------------------
        # 3. Kiểm tra đây có thực sự là file không
        # ----------------------------------------------------

        if not path.is_file():
            raise ValueError(
                f"Path is not a file: {path}"
            )


        # ----------------------------------------------------
        # 4. Lấy extension
        # ----------------------------------------------------
        #
        # booking_policy.md
        #                 ↓
        #                ".md"
        #
        # lower() giúp:
        #
        # POLICY.PDF
        #
        # vẫn trở thành:
        #
        # ".pdf"
        # ----------------------------------------------------

        extension = path.suffix.lower()


        # ----------------------------------------------------
        # 5. Kiểm tra loại file có được hỗ trợ không
        # ----------------------------------------------------

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {extension}"
            )


        # ====================================================
        # SỬA CHỖ 2:
        #
        # Trước đây bạn có:
        #
        # text = path.read_text(encoding="utf-8")
        #
        # Dòng đó chỉ phù hợp với text file:
        #
        # .txt
        # .md
        #
        # PDF và DOCX không phải plain text,
        # nên phải chọn parser dựa trên extension.
        #
        # Đây gọi là dispatch theo file type.
        # ====================================================

        if extension in {".txt", ".md"}:

            # File text thuần
            text = self._load_text_file(path)

        elif extension == ".pdf":

            # PDF dùng pypdf
            text = self._load_pdf(path)

        elif extension == ".docx":

            # Word .docx dùng python-docx
            text = self._load_docx(path)

        else:

            # Về lý thuyết code này không chạy vì phía trên
            # đã check SUPPORTED_EXTENSIONS.
            #
            # Nhưng vẫn giữ defensive programming.
            raise ValueError(
                f"Unsupported file type: {extension}"
            )


        # ----------------------------------------------------
        # 6. Kiểm tra file có nội dung hay không
        # ----------------------------------------------------
        #
        # Điều này đặc biệt hữu ích với PDF.
        #
        # Một file PDF scan có thể chứa toàn ảnh,
        # pypdf extract ra text rỗng.
        #
        # Khi đó text.strip() == ""
        # ----------------------------------------------------

        if not text.strip():
            raise ValueError(
                f"Knowledge file is empty or contains no extractable text: {path}"
            )


        # ----------------------------------------------------
        # 7. Tạo Document
        # ----------------------------------------------------

        return Document(
            text=text,
            source=path.name,
            file_path=str(path),
        )


    # ========================================================
    # SỬA CHỖ 3:
    # Tách logic đọc text file thành method riêng.
    # ========================================================
    #
    # Tại sao không để trực tiếp trong load_file()?
    #
    # Vì sau này load_file chỉ nên chịu trách nhiệm:
    #
    # extension
    #      ↓
    # chọn parser
    #
    # Còn từng parser xử lý việc đọc riêng.
    #
    # Điều này giúp code dễ mở rộng.
    # ========================================================

    def _load_text_file(
        self,
        path: Path,
    ) -> str:
        """
        Đọc file text thuần như .txt và .md.
        """

        return path.read_text(
            encoding="utf-8"
        )


    # ========================================================
    # SỬA CHỖ 4:
    # Thêm parser dành cho PDF.
    # ========================================================

    def _load_pdf(
        self,
        path: Path,
    ) -> str:
        """
        Extract text từ PDF.

        Flow:

        PDF
         ↓
        PdfReader
         ↓
        từng page
         ↓
        extract_text()
         ↓
        ghép lại thành một string
        """

        # PdfReader đọc cấu trúc file PDF.
        reader = PdfReader(path)

        # Danh sách dùng để chứa text của từng page.
        pages: list[str] = []


        # ----------------------------------------------------
        # Duyệt từng page trong PDF
        # ----------------------------------------------------

        for page in reader.pages:

            # extract_text() cố lấy text layer trong trang PDF.
            #
            # Nếu PDF là file export từ Word / browser / report
            # thì thường sẽ lấy được text.
            #
            # Nếu PDF là ảnh scan thì có thể trả về None hoặc "".
            page_text = page.extract_text()


            # Chỉ append nếu page có text.
            if page_text and page_text.strip():

                pages.append(
                    page_text.strip()
                )


        # ----------------------------------------------------
        # Ghép các page lại.
        # ----------------------------------------------------
        #
        # Dùng "\n\n" để giữ khoảng cách giữa các page.
        #
        # Ví dụ:
        #
        # page 1 text
        #
        # page 2 text
        #
        # page 3 text
        # ----------------------------------------------------

        return "\n\n".join(pages)


    # ========================================================
    # THÊM MỚI:
    # Parser dành cho Word .docx.
    # ========================================================

    def _load_docx(
        self,
        path: Path,
    ) -> str:
        """
        Extract text từ file Word .docx.

        Flow:

        DOCX
         ↓
        python-docx
         ↓
        paragraphs
         ↓
        lấy paragraph.text
         ↓
        ghép lại thành string
        """

        # DocxDocument ở đây là class Document của python-docx
        # mà chúng ta đã alias ở phần import.
        document = DocxDocument(path)


        # Chứa text của từng paragraph.
        paragraphs: list[str] = []


        # ----------------------------------------------------
        # Duyệt từng paragraph trong Word
        # ----------------------------------------------------

        for paragraph in document.paragraphs:

            # paragraph.text trả về text của paragraph.
            #
            # strip() loại whitespace dư đầu/cuối.
            paragraph_text = paragraph.text.strip()


            # Bỏ qua paragraph rỗng.
            if paragraph_text:

                paragraphs.append(
                    paragraph_text
                )


        # ----------------------------------------------------
        # Ghép paragraph thành text hoàn chỉnh.
        # ----------------------------------------------------
        #
        # Mỗi paragraph cách nhau bằng 2 newline.
        #
        # Điều này hữu ích cho chunker sau này vì vẫn giữ được
        # một phần cấu trúc paragraph của document.
        # ----------------------------------------------------

        return "\n\n".join(paragraphs)


    def load_directory(
        self,
        directory_path: str | Path,
    ) -> list[Document]:
        """
        Load tất cả file knowledge được hỗ trợ trong một folder.

        Ví dụ:

        knowledge/
            ├── booking.md
            ├── faq.txt
            ├── policy.pdf
            └── handbook.docx

                  ↓

        load_directory()

                  ↓

        [
            Document(...),
            Document(...),
            Document(...),
            Document(...)
        ]
        """

        directory = Path(directory_path)


        # ----------------------------------------------------
        # 1. Kiểm tra directory tồn tại
        # ----------------------------------------------------

        if not directory.exists():
            raise FileNotFoundError(
                f"Knowledge directory does not exist: {directory}"
            )


        # ----------------------------------------------------
        # 2. Kiểm tra đây có phải folder không
        # ----------------------------------------------------

        if not directory.is_dir():
            raise ValueError(
                f"Path is not a directory: {directory}"
            )


        # List chứa tất cả Document sau khi load.
        documents: list[Document] = []


        # ----------------------------------------------------
        # 3. Duyệt từng item trong folder
        # ----------------------------------------------------
        #
        # sorted() giúp thứ tự load deterministic.
        #
        # Nghĩa là chạy nhiều lần vẫn cùng thứ tự,
        # rất hữu ích khi debug/indexing.
        # ----------------------------------------------------

        for path in sorted(directory.iterdir()):

            # Nếu là folder con thì hiện tại bỏ qua.
            #
            # Ví dụ:
            #
            # knowledge/
            #   ├── a.pdf
            #   └── policies/
            #       └── b.pdf
            #
            # b.pdf sẽ CHƯA được load.
            #
            # Sau này muốn recursive có thể dùng rglob().
            if not path.is_file():
                continue


            # ------------------------------------------------
            # Chỉ xử lý extension được hỗ trợ.
            # ------------------------------------------------
            #
            # Vì SUPPORTED_EXTENSIONS giờ đã bao gồm:
            #
            # .txt
            # .md
            # .pdf
            # .docx
            #
            # nên load_directory KHÔNG cần sửa logic gì thêm.
            # ------------------------------------------------

            if path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
                continue


            # ------------------------------------------------
            # 4. Tái sử dụng load_file()
            # ------------------------------------------------
            #
            # Đây là điểm kiến trúc của bạn đang làm đúng.
            #
            # load_directory không cần biết:
            #
            # PDF đọc thế nào?
            # DOCX đọc thế nào?
            #
            # Nó chỉ gọi:
            #
            # self.load_file(path)
            #
            # Còn load_file tự dispatch parser.
            # ------------------------------------------------

            document = self.load_file(path)

            documents.append(document)


        return documents